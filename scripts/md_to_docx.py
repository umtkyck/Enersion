#!/usr/bin/env python3
"""
Markdown to DOCX Converter for Enersion Project
Converts PROJECT_STATUS.md to a formatted Word document
"""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_docx_from_markdown(md_file, output_file):
    """Convert markdown file to DOCX with formatting"""
    
    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set document properties
    core_props = doc.core_properties
    core_props.title = "Enersion Linux Application - Project Status"
    core_props.author = "Enersion Team"
    core_props.subject = "Project Status Report"
    
    # Define styles
    styles = doc.styles
    
    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 102, 153)
    
    # Heading 3 style
    h3_style = styles['Heading 3']
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(51, 51, 51)
    
    # Process markdown line by line
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_content = []
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                if code_content:
                    p = doc.add_paragraph()
                    p.style = 'No Spacing'
                    for code_line in code_content:
                        run = p.add_run(code_line + '\n')
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    # Add border/shading effect
                    p.paragraph_format.left_indent = Inches(0.3)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        # Handle tables
        if '|' in line and not line.strip().startswith('```'):
            # Check if it's a table separator line
            if re.match(r'^\s*\|[\s\-:|]+\|\s*$', line):
                i += 1
                continue
            
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                if not in_table:
                    in_table = True
                    table_rows = []
                table_rows.append(cells)
            i += 1
            
            # Check if next line is not a table
            if i >= len(lines) or ('|' not in lines[i] or lines[i].strip().startswith('```')):
                # Create table
                if table_rows:
                    num_cols = max(len(row) for row in table_rows)
                    table = doc.add_table(rows=len(table_rows), cols=num_cols)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    for row_idx, row_data in enumerate(table_rows):
                        row = table.rows[row_idx]
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                cell.text = cell_data.replace('✅', '[OK]').replace('🔄', '[WIP]').replace('📋', '[TODO]')
                                
                                # Header row styling
                                if row_idx == 0:
                                    set_cell_shading(cell, '003366')
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                                            run.font.color.rgb = RGBColor(255, 255, 255)
                                            run.font.size = Pt(10)
                                else:
                                    # Alternate row colors
                                    if row_idx % 2 == 0:
                                        set_cell_shading(cell, 'F0F0F0')
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.size = Pt(10)
                    
                    doc.add_paragraph()  # Space after table
                    table_rows = []
                in_table = False
            continue
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            # Title/H1
            text = line[2:].strip()
            text = re.sub(r'[📊🎯🔧📁✅🔄📋🐛📈📝📞]', '', text).strip()
            p = doc.add_heading(text, level=0)
            i += 1
            continue
        
        if line.startswith('## '):
            text = line[3:].strip()
            text = re.sub(r'[📊🎯🔧📁✅🔄📋🐛📈📝📞]', '', text).strip()
            doc.add_heading(text, level=1)
            i += 1
            continue
        
        if line.startswith('### '):
            text = line[4:].strip()
            text = re.sub(r'[📊🎯🔧📁✅🔄📋🐛📈📝📞]', '', text).strip()
            doc.add_heading(text, level=2)
            i += 1
            continue
        
        if line.startswith('#### '):
            text = line[5:].strip()
            text = re.sub(r'[📊🎯🔧📁✅🔄📋🐛📈📝📞]', '', text).strip()
            doc.add_heading(text, level=3)
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.add_run('─' * 60)
            i += 1
            continue
        
        # Handle bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Handle checkboxes
            text = text.replace('[x]', '☑').replace('[ ]', '☐')
            text = re.sub(r'[📊🎯🔧📁✅🔄📋🐛📈📝📞]', '', text).strip()
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        
        # Handle numbered lists
        match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if match:
            text = match.group(2)
            text = text.replace('[x]', '☑').replace('[ ]', '☐')
            text = re.sub(r'[📊🎯🔧📁✅🔄📋🐛📈📝📞]', '', text).strip()
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Regular paragraph
        text = line.strip()
        if text:
            # Remove emojis
            text = re.sub(r'[📊🎯🔧📁✅🔄📋🐛📈📝📞⚡💾💻🔧🐍🖥️📖]', '', text).strip()
            # Handle bold text
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            # Handle inline code
            text = re.sub(r'`(.+?)`', r'[\1]', text)
            
            if text:
                p = doc.add_paragraph(text)
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f"Document saved to: {output_file}")

def main():
    # Paths
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)
    
    md_file = os.path.join(project_root, 'LinuxApp', 'PROJECT_STATUS.md')
    output_file = os.path.join(project_root, 'Enersion_Project_Status.docx')
    
    if os.path.exists(md_file):
        create_docx_from_markdown(md_file, output_file)
    else:
        print(f"Error: {md_file} not found")

if __name__ == '__main__':
    main()

